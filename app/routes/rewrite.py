"""Rewrite route — AI-powered rewriting endpoints."""

from __future__ import annotations

import io
from typing import Literal

from fastapi import APIRouter, Depends, HTTPException, Request, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from app.dependencies.rate_limit import enforce_usage_limit, record_usage
from app.tools.rewriter_tool import rewrite_paragraph, rewrite_document
from app.tools.general_rewriter import general_rewrite
from app.utils.logger import get_logger

logger = get_logger(__name__)

router = APIRouter(prefix="/api/v1", tags=["rewrite"])


# ---------------------------------------------------------------------------
# Request / response schemas
# ---------------------------------------------------------------------------

class RewriteParagraphRequest(BaseModel):
    """Request body for paragraph rewriting."""

    text: str = Field(..., min_length=1, description="The flagged paragraph to rewrite")
    context: str = Field(default="", description="Surrounding text for context")
    tone: str = Field(default="academic", description="Writing tone: academic, professional, casual")


class RewriteParagraphResponse(BaseModel):
    """Response for paragraph rewriting."""

    original: str
    rewritten: str
    tone: str
    elapsed_s: float
    skipped: bool = False
    skip_reason: str = ""


class RewriteDocumentRequest(BaseModel):
    """Request body for full document rewriting."""

    text: str = Field(..., min_length=1, description="The full document text")
    flagged_passages: list[str] = Field(
        default_factory=list,
        description="Passages flagged as plagiarised (will be targeted for rewriting)",
    )
    tone: str = Field(default="academic", description="Writing tone: academic, professional, casual")


class RewriteDocumentResponse(BaseModel):
    """Response for document rewriting."""

    original: str
    rewritten: str
    passages_rewritten: int
    elapsed_s: float


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------

@router.post(
    "/rewrite/paragraph",
    response_model=RewriteParagraphResponse,
    status_code=status.HTTP_200_OK,
    summary="Rewrite a single paragraph to eliminate plagiarism",
    dependencies=[Depends(enforce_usage_limit)],
)
async def rewrite_paragraph_endpoint(
    request: RewriteParagraphRequest,
    http_request: Request = None,
) -> RewriteParagraphResponse:
    """Accept a flagged paragraph, rewrite it using AI, and return the
    rewritten version with the original preserved for comparison."""
    logger.info(
        "rewrite_paragraph_requested",
        text_length=len(request.text),
        tone=request.tone,
    )

    try:
        result = await rewrite_paragraph(
            text=request.text,
            context=request.context,
            tone=request.tone,
        )
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(exc),
        )
    except RuntimeError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"AI rewrite failed: {exc}",
        )

    # Record usage
    if http_request:
        record_usage(http_request, tool_type="rewrite")

    # Bridge tool output (rewrites list) to response model (rewritten string)
    rewrites = result.get("rewrites", [])
    best = rewrites[0] if rewrites else result["original"]
    # Guard: if AI returned a dict instead of a string, extract text
    if isinstance(best, dict):
        best = best.get("text") or best.get("rewrite") or best.get("rewritten") or str(best)
    return RewriteParagraphResponse(
        original=result["original"],
        rewritten=str(best),
        tone=result.get("tone", request.tone),
        elapsed_s=result.get("elapsed_s", 0.0),
        skipped=result.get("skipped", False),
        skip_reason=result.get("skip_reason", ""),
    )


@router.post(
    "/rewrite/document",
    response_model=RewriteDocumentResponse,
    status_code=status.HTTP_200_OK,
    summary="Rewrite flagged passages in a full document to eliminate plagiarism",
    dependencies=[Depends(enforce_usage_limit)],
)
async def rewrite_document_endpoint(
    request: RewriteDocumentRequest,
    http_request: Request = None,
) -> RewriteDocumentResponse:
    """Accept a full document and a list of flagged passages, rewrite
    only the flagged sections using AI, and return the complete document
    with rewrites applied."""
    logger.info(
        "rewrite_document_requested",
        doc_length=len(request.text),
        flagged_count=len(request.flagged_passages),
        tone=request.tone,
    )

    try:
        result = await rewrite_document(
            document_text=request.text,
            flagged_passages=request.flagged_passages,
            tone=request.tone,
        )
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(exc),
        )
    except RuntimeError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"AI rewrite failed: {exc}",
        )

    # Record usage
    if http_request:
        record_usage(http_request, tool_type="rewrite")

    return RewriteDocumentResponse(**result)


# ---------------------------------------------------------------------------
# General-purpose rewriter (NEW — not tied to plagiarism)
# ---------------------------------------------------------------------------

class GeneralRewriteRequest(BaseModel):
    """Request body for the general-purpose rewriting tool."""

    text: str = Field(..., min_length=1, description="Text to rewrite")
    mode: Literal[
        "paraphrase", "simplify", "expand",
        "formal", "casual", "academic", "humanize",
    ] = Field(default="paraphrase", description="Rewriting mode")
    tone: Literal[
        "friendly", "professional", "confident",
        "persuasive", "neutral",
    ] = Field(default="neutral", description="Desired tone")
    strength: Literal["low", "medium", "high"] = Field(
        default="medium", description="Rewrite intensity"
    )


class GeneralRewriteResponse(BaseModel):
    """Response for the general-purpose rewriting tool."""

    original: str
    variations: list[str]
    mode: str
    tone: str
    strength: str
    skipped: bool = False
    skip_reason: str = ""
    elapsed_s: float


@router.post(
    "/rewrite/general",
    response_model=GeneralRewriteResponse,
    status_code=status.HTTP_200_OK,
    summary="Rewrite text using the general-purpose writing assistant",
    dependencies=[Depends(enforce_usage_limit)],
)
async def general_rewrite_endpoint(
    request: GeneralRewriteRequest,
    http_request: Request = None,
) -> GeneralRewriteResponse:
    """Rewrite any text with mode, tone, and strength controls.

    Unlike the plagiarism rewriter, this is a general writing assistant
    that produces 3 diverse variations for every request.
    """
    logger.info(
        "general_rewrite_requested",
        text_length=len(request.text),
        mode=request.mode,
        tone=request.tone,
        strength=request.strength,
    )

    try:
        result = await general_rewrite(
            text=request.text,
            mode=request.mode,
            tone=request.tone,
            strength=request.strength,
        )
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(exc),
        )
    except RuntimeError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"AI rewrite failed: {exc}",
        )

    # Record usage
    if http_request:
        record_usage(http_request, tool_type="rewrite")

    return GeneralRewriteResponse(**result)


# ---------------------------------------------------------------------------
# Export rewritten text as Word (.docx) with tracked changes
# ---------------------------------------------------------------------------

class ExportDocxRequest(BaseModel):
    """Request body for .docx export."""

    original: str = Field(..., min_length=1, description="Original text")
    rewritten: str = Field(..., min_length=1, description="Rewritten text")
    title: str = Field(default="Rewritten Document", max_length=200)
    show_changes: bool = Field(default=True, description="Highlight changed sections")


@router.post(
    "/rewrite/export-docx",
    summary="Export rewritten text as a Word document",
    responses={200: {"content": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document": {}}}},
)
async def export_rewrite_docx(request: ExportDocxRequest):
    """Generate a .docx file with the rewritten text.

    When ``show_changes`` is True, the document contains both the original
    (struck through in red) and the rewritten text (in green) side by side
    for easy comparison.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_COLOR_INDEX

    doc = Document()

    # Title
    title_para = doc.add_heading(request.title, level=1)

    if request.show_changes:
        # Split into paragraphs and show diff
        orig_paras = request.original.strip().split("\n")
        new_paras = request.rewritten.strip().split("\n")

        doc.add_heading("Rewritten Document (Changes Highlighted)", level=2)
        doc.add_paragraph(
            "Red strikethrough = removed text. Green = new text.",
            style="Intense Quote" if "Intense Quote" in [s.name for s in doc.styles] else None,
        )

        max_len = max(len(orig_paras), len(new_paras))
        for i in range(max_len):
            orig = orig_paras[i].strip() if i < len(orig_paras) else ""
            new = new_paras[i].strip() if i < len(new_paras) else ""

            para = doc.add_paragraph()

            if orig == new:
                # Unchanged
                run = para.add_run(new)
                run.font.size = Pt(11)
            else:
                if orig:
                    # Show removed text
                    run_old = para.add_run(orig + " ")
                    run_old.font.size = Pt(11)
                    run_old.font.color.rgb = RGBColor(0xE1, 0x70, 0x55)
                    run_old.font.strike = True
                if new:
                    # Show new text
                    run_new = para.add_run(new)
                    run_new.font.size = Pt(11)
                    run_new.font.color.rgb = RGBColor(0x00, 0xB8, 0x94)
    else:
        # Just the clean rewritten text
        doc.add_heading("Rewritten Document", level=2)
        for para_text in request.rewritten.strip().split("\n"):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                for run in p.runs:
                    run.font.size = Pt(11)

    # Add footer note
    doc.add_paragraph("")
    footer = doc.add_paragraph("Generated by PlagiarismGuard — plagiarismguard.com")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0x8B, 0x8F, 0xA3)

    # Write to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = request.title.replace(" ", "_")[:50] + ".docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
